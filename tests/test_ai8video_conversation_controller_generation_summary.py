from __future__ import annotations

from pathlib import Path
import tempfile
import unittest
from unittest.mock import Mock
from unittest.mock import patch

from ai8video.application.conversation_controller import AI8VideoConversationController
from ai8video.core.models import ConversationState, VideoPrompt, ParsedRequest, PipelineResult, QuickVideoJob, GenerationOutcome
from ai8video.assets import user_materials


class AI8VideoConversationControllerGenerationSummaryTest(unittest.TestCase):
    def setUp(self) -> None:
        self.default_script_reference_patcher = patch(
            "ai8video.knowledge.default_script_reference.load_default_script_reference",
            return_value=None,
        )
        self.default_reference_image_patcher = patch(
            "ai8video.application.conversation_controller.default_reference_image_path",
            return_value=None,
        )
        self.default_generation_mode_patcher = patch(
            "ai8video.application.conversation_controller.default_concurrent_generation_enabled",
            return_value=False,
        )
        self.default_html_motion_patcher = patch(
            "ai8video.application.conversation_controller.default_html_motion_overlay_enabled",
            return_value=False,
        )
        self.default_script_reference_patcher.start()
        self.default_reference_image_patcher.start()
        self.default_generation_mode_patcher.start()
        self.default_html_motion_patcher.start()

    def tearDown(self) -> None:
        self.default_generation_mode_patcher.stop()
        self.default_html_motion_patcher.stop()
        self.default_reference_image_patcher.stop()
        self.default_script_reference_patcher.stop()

    @staticmethod
    def _agent(pipeline: Mock) -> AI8VideoConversationController:
        return AI8VideoConversationController(pipeline, merge_mode_loader=lambda: "none")

    @staticmethod
    def _build_result(*, generated: bool = True) -> PipelineResult:
        request = ParsedRequest(raw_text="负责人在会议室讲素材返工风险", mode="single_video")
        video = VideoPrompt(index=1, title="单条视频", prompt="负责人在会议室讲素材返工风险")
        job = QuickVideoJob(
            video_index=1,
            job_id="job-generation-summary",
            status="succeeded" if generated else "failed",
            prompt=video.prompt,
            storage_key="mobile:job-generation-summary" if generated else None,
            final_frame_storage_key="mobile:job-generation-summary/final" if generated else None,
        )
        outcome = GenerationOutcome(
            video_index=1,
            job_id=job.job_id,
            status=job.status,
            decision="generated" if generated else "failed",
            reasons=[] if generated else ["生成失败"],
            meta={"kind": "generation_outcome"},
        )
        return PipelineResult(
            request=request,
            videos=[video],
            first_frame=None,
            jobs=[job],
            outcomes=[outcome],
            archives=[],
            asset_records=[],
            dry_run=True,
        )

    def test_completed_reply_mentions_generation_and_archive_flow(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message("generation-summary", "生成一条短视频：负责人在会议室讲素材返工风险")

        self.assertEqual(reply.stage, "completed")
        self.assertIn("创建任务和归档", reply.text)
        request = pipeline.run_request.call_args.args[0]
        self.assertIsNone(request.reference_image)

    def test_completed_reply_does_not_mention_removed_review_chain(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message("generation-clean", "生成一条短视频：负责人在会议室讲素材返工风险")

        self.assertNotIn("审" + "核", reply.text)
        self.assertNotIn("抽" + "检", reply.text)

    def test_explicit_reference_request_still_waits_when_tab_has_no_selection(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message(
            "generation-explicit-reference",
            "生成一条短视频：负责人在会议室讲素材返工风险，需要参考图",
        )

        self.assertEqual(reply.awaiting, "reference_image")
        self.assertIn("标签页当前没有选中图片", reply.text)
        pipeline.run_request.assert_not_called()

    def test_multi_video_request_with_count_and_reference_generates_when_keywords_present(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message(
            "generation-direct-run",
            "根据这个剧本生成 2 个 10s 短视频，老板商务风。核心主题：全球发布倒计时。参考图：/tmp/612.png",
        )

        self.assertEqual(reply.stage, "completed")
        pipeline.run_request.assert_called_once()
        request = pipeline.run_request.call_args.args[0]
        self.assertEqual(request.video_count, 2)
        self.assertEqual(request.duration_seconds, 10)
        self.assertTrue(request.iterative_generation)
        self.assertFalse(request.concurrent_generation)
        self.assertEqual(request.reference_image, "/tmp/612.png")
        self.assertEqual(request.core_keywords, "全球发布倒计时")

    def test_multi_video_generation_forces_serial_iteration_even_when_concurrency_is_requested(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message(
            "generation-concurrent-choice",
            "根据这个剧本生成 2 个 10s 短视频，老板商务风。核心主题：6月18日全球发布倒计时、私域资产。参考图：/tmp/612.png，并发模式",
        )

        self.assertEqual(reply.stage, "completed")
        pipeline.run_request.assert_called_once()
        request = pipeline.run_request.call_args.args[0]
        self.assertFalse(request.concurrent_generation)
        self.assertTrue(request.iterative_generation)
        self.assertEqual(request.core_keywords, "6月18日全球发布倒计时、私域资产")

    def test_six_video_request_is_rejected_before_pipeline_submission(self) -> None:
        pipeline = Mock()
        agent = self._agent(pipeline)

        reply = agent.handle_message(
            "generation-over-limit",
            "生成 6 个 10s 短视频，老板商务风。核心主题：私域资产。参考图：/tmp/612.png",
        )

        self.assertEqual(reply.stage, "collecting")
        self.assertEqual(reply.awaiting, "video_count")
        self.assertEqual(reply.meta["validation"], "iterative_batch_limit")
        self.assertIn("最多生成 5 条", reply.text)
        pipeline.run_request.assert_not_called()

    def test_html_motion_setting_is_snapshotted_into_generation_request(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        with patch(
            "ai8video.application.conversation_controller.default_html_motion_overlay_enabled",
            return_value=True,
        ):
            reply = agent.handle_message(
                "generation-html-motion",
                "根据这个剧本生成 2 个 10s 短视频，老板商务风。核心主题：私域资产。参考图：/tmp/612.png",
            )

        self.assertEqual(reply.stage, "completed")
        request = pipeline.run_request.call_args.args[0]
        self.assertTrue(request.html_motion_overlay_enabled)

    def test_multi_video_defaults_to_normal_generation_without_mode_prompt(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message(
            "generation-normal-choice",
            "根据这个剧本生成 2 个 10s 短视频，老板商务风。核心主题：私域资产。参考图：/tmp/612.png",
        )

        self.assertEqual(reply.stage, "completed")
        pipeline.run_request.assert_called_once()
        request = pipeline.run_request.call_args.args[0]
        self.assertFalse(request.concurrent_generation)
        self.assertEqual(request.core_keywords, "私域资产")

    def test_plain_number_core_keyword_followup_uses_auto_extraction(self) -> None:
        agent = self._agent(Mock())
        state = ConversationState(session_id="numeric-core-keywords", awaiting="core_keywords")

        agent._handle_core_keywords_followup(state, "30")

        self.assertEqual(
            state.draft.core_keywords,
            "按用户提供的原文自行提炼核心主题，但不得偏离用户原始要求",
        )
        self.assertIsNone(state.awaiting)

    def test_docx_material_is_read_reported_and_passed_to_generation(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as tempdir:
            root = Path(tempdir)
            image_dir = root / "图片素材库"
            script_dir = root / "剧本素材库"
            image_dir.mkdir()
            script_dir.mkdir()
            docx_path = script_dir / "2.docx"
            document = Document()
            document.add_paragraph("这一批我建议直接围绕AI8video 全球发布倒计时打造全网热点。")
            document.add_paragraph("核心策略不是介绍 App，而是制造行业关注。")
            document.save(str(docx_path))

            pipeline = Mock()
            pipeline.run_request.return_value = self._build_result()
            agent = self._agent(pipeline)

            with patch.object(user_materials, "USER_MATERIAL_ROOT", root), patch.object(
                user_materials, "USER_IMAGE_MATERIAL_DIR", image_dir
            ), patch.object(user_materials, "USER_SCRIPT_MATERIAL_DIR", script_dir):
                reply = agent.handle_message(
                    "generation-docx-material",
                    "@2.docx 生成 2 个 10s 短视频。参考图：/tmp/612.png",
                )

            self.assertEqual(reply.stage, "completed")
            self.assertIn("已读取剧本素材：@2.docx", reply.text)
            self.assertIn("正文约", reply.text)
            self.assertIn("这一批我建议直接围绕AI8video", reply.text)
            pipeline.run_request.assert_called_once()
            request = pipeline.run_request.call_args.args[0]
            self.assertIn("@2.docx 剧本素材内容", request.raw_text)
            self.assertIn("核心策略不是介绍 App", request.raw_text)
            self.assertFalse(request.concurrent_generation)
            self.assertIsNone(request.core_keywords)

    def test_missing_dialogue_prompts_auto_uses_smart_completion_before_generation(self) -> None:
        pipeline = Mock()
        pipeline.run_request.return_value = self._build_result()
        agent = self._agent(pipeline)

        reply = agent.handle_message(
            "generation-missing-dialogue",
            "根据 /tmp/612.png 生成 3 个 10s 短视频，老板商务风。",
        )

        self.assertEqual(reply.stage, "completed")
        self.assertIsNone(reply.awaiting)
        self.assertNotIn("guide", reply.meta)
        pipeline.run_request.assert_called_once()
        request = pipeline.run_request.call_args.args[0]
        self.assertIn("口播文案补全要求", request.raw_text)
        self.assertFalse(request.concurrent_generation)
        self.assertIsNone(request.core_keywords)


if __name__ == "__main__":
    unittest.main()
